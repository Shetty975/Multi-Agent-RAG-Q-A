# import requests
# import logging
# import re
# from retriever import retrieve_top_k
# from langchain.chat_models import ChatOpenAI
# from langchain.chains.question_answering import load_qa_chain
# from groq import Groq

# # --------------------------------------------
# # Logging setup
# # --------------------------------------------
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s', filename='agent.log')

# # --------------------------------------------
# # API KEYS
# # --------------------------------------------
# WOLFRAM_APP_ID = "VA22W6-UJVQHGQUYE"
# GROQ_API_KEY = "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"
# APP_ID = "6668edc7"
# APP_KEY = "1ca66b8f5b37fe249111bca5eb18d9ac"

# # --------------------------------------------
# # Clients
# # --------------------------------------------
# groq_client = Groq(api_key=GROQ_API_KEY)

# # --------------------------------------------
# # Math Solvers
# # --------------------------------------------
# def solve_math_problem_with_wolfram(query):
#     api_endpoint = "http://api.wolframalpha.com/v2/query"
#     params = {
#         'input': query,
#         'output': 'json',
#         'appid': WOLFRAM_APP_ID
#     }

#     try:
#         response = requests.get(api_endpoint, params=params)
#         response.raise_for_status()
#         data = response.json()
#         pods = data.get('queryresult', {}).get('pods', [])

#         if pods:
#             for pod in pods:
#                 if pod.get('title', '').lower() in ['result', 'exact result']:
#                     for subpod in pod.get('subpods', []):
#                         if subpod.get('plaintext'):
#                             return subpod['plaintext']
#             for pod in pods:
#                 for subpod in pod.get('subpods', []):
#                     if subpod.get('plaintext'):
#                         return subpod['plaintext']
#         return "No result found."
#     except Exception as e:
#         return f"Error querying Wolfram Alpha: {e}"

# def solve_math_with_groq_llm(query):
#     prompt = f"Solve this mathematical problem with detailed steps: {query}"
#     try:
#         response = groq_client.chat.completions.create(
#             model="mixtral-8x7b-32768",
#             messages=[{"role": "user", "content": prompt}],
#             temperature=0
#         )
#         return response.choices[0].message.content.strip()
#     except Exception as e:
#         return f"Groq LLM error: {e}"

# # --------------------------------------------
# # Free Dictionary API
# # --------------------------------------------
# FREE_DICT_BASE_URL = "https://api.dictionaryapi.dev/api/v2/entries/en/"

# def get_definition_from_free_dictionary(query):
#     url = f"{FREE_DICT_BASE_URL}{query.lower()}"

#     try:
#         response = requests.get(url)
#         response.raise_for_status()
#         data = response.json()

#         word = data[0]['word']
#         definition = data[0]['meanings'][0]['definitions'][0]['definition']
#         return word, definition
#     except Exception as e:
#         return None, f"Free Dictionary API error: {str(e)}"

# # --------------------------------------------
# # Logging Helper
# # --------------------------------------------
# def log_decision(query, route):
#     logging.info(f"Query: '{query}' | Route: {route}")

# # --------------------------------------------
# # LLM Answer (LangChain)
# # --------------------------------------------
# def get_llm_answer(query, docs):
#     llm = ChatOpenAI(temperature=0)
#     chain = load_qa_chain(llm, chain_type="stuff")
#     return chain.run(input_documents=docs, question=query)

# # --------------------------------------------
# # Query Classifiers
# # --------------------------------------------
# def is_math_query(query):
#     return re.search(r'\b(calculate|plus|minus|times|divide|compute|multiplied|divided|add|subtract|\d+|\=|\^|\*|/|-|\+)\b', query.lower())

# def is_definition_query(query):
#     return re.search(r'\b(define|what is|meaning of|definition of|explain|describe)\b', query.lower())

# # --------------------------------------------
# # Agent Router
# # --------------------------------------------
# def agent_router(query):
#     response = {}

#     if is_math_query(query):
#         try:
#             solution = solve_math_problem_with_wolfram(query)
#             if "No result found" in solution or "Error" in solution:
#                 solution = solve_math_with_groq_llm(query)  # Fallback
#             response['route'] = '🧮 Calculator (Wolfram → Groq Fallback)'
#             response['result'] = solution
#         except Exception:
#             response['route'] = '🧮 Calculator (Groq Only)'
#             response['result'] = solve_math_with_groq_llm(query)
#         finally:
#             log_decision(query, response['route'])

#     elif is_definition_query(query):
#         try:
#             word = re.sub(r'^(define|what is|meaning of|definition of|explain|describe)\s*', '', query.lower()).strip()
#             word, definition = get_definition_from_free_dictionary(word)
#             if word and definition:
#                 response['route'] = '📚 Dictionary (Free Dictionary API)'
#                 response['result'] = f"Definition of {word}: {definition}"
#             else:
#                 response['route'] = '📚 Dictionary (Not Found)'
#                 response['result'] = "No definition found."
#         except Exception as e:
#             response['route'] = '📚 Dictionary (Error)'
#             response['result'] = f"Definition error: {e}"
#         finally:
#             log_decision(query, response['route'])

#     else:
#         try:
#             docs = retrieve_top_k(query)
#             context = [doc.page_content[:200] for doc in docs]
#             answer = get_llm_answer(query, docs)
#             response['route'] = '📖 RAG → LLM'
#             response['context'] = context
#             response['result'] = answer
#         except Exception as e:
#             response['route'] = '📖 RAG → LLM (Error)'
#             response['result'] = f"RAG error: {e}"
#         finally:
#             log_decision(query, response['route'])

#     return response

# # --------------------------------------------
# # RAG: Document Ingestion & Retrieval
# # --------------------------------------------
# import os
# import faiss
# from sentence_transformers import SentenceTransformer
# from transformers import pipeline
# from langchain.schema import Document
# from PyPDF2 import PdfReader
# import docx
# import pandas as pd

# # Embedding and QA models
# embed_model = SentenceTransformer('all-MiniLM-L6-v2')
# qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2")

# def chunk_text(text, max_length=300):
#     sentences = re.split(r'[.!?]\s+', text)
#     chunks, chunk = [], ""
#     for sentence in sentences:
#         if len(chunk) + len(sentence) < max_length:
#             chunk += sentence + '. '
#         else:
#             chunks.append(chunk.strip())
#             chunk = sentence + '. '
#     if chunk:
#         chunks.append(chunk.strip())
#     return chunks

# def read_file(file):
#     ext = os.path.splitext(file.name)[-1].lower()
#     if ext == '.pdf':
#         reader = PdfReader(file)
#         return "\n".join([page.extract_text() or "" for page in reader.pages])
#     elif ext == '.docx':
#         doc = docx.Document(file)
#         return "\n".join([p.text for p in doc.paragraphs])
#     elif ext == '.csv':
#         df = pd.read_csv(file)
#         return df.to_string(index=False)
#     elif ext in ['.txt', '.md']:
#         return file.read().decode("utf-8")
#     else:
#         return ""

# def process_documents(files):
#     texts, metadatas = [], []
#     for file in files:
#         content = read_file(file)
#         if content:
#             chunks = chunk_text(content)
#             texts.extend(chunks)
#             metadatas.extend([{"source": file.name}] * len(chunks))
#     return texts, metadatas

# def build_vector_store(texts):
#     embeddings = embed_model.encode(texts)
#     index = faiss.IndexFlatL2(embeddings[0].shape[0])
#     index.add(embeddings)
#     return index, embeddings

# def retrieve_top_k_local(query, texts, index, k=3):
#     query_vec = embed_model.encode([query])
#     _, indices = index.search(query_vec, k)
#     return [Document(page_content=texts[i]) for i in indices[0]]



import os
import re
import faiss
import logging
import requests
import docx
import pandas as pd

from PyPDF2 import PdfReader
from groq import Groq
from sentence_transformers import SentenceTransformer
from transformers import pipeline
from langchain.schema import Document
from langchain.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain

# --------------------------------------------
# Logging setup
# --------------------------------------------
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s', filename='agent.log')

# --------------------------------------------
# API Keys
# --------------------------------------------
WOLFRAM_APP_ID = "VA22W6-UJVQHGQUYE"
GROQ_API_KEY = "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"
APP_ID = "6668edc7"
APP_KEY = "1ca66b8f5b37fe249111bca5eb18d9ac"

# --------------------------------------------
# Clients
# --------------------------------------------
groq_client = Groq(api_key=GROQ_API_KEY)
embed_model = SentenceTransformer('all-MiniLM-L6-v2')
qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2")

# --------------------------------------------
# Math Solvers
# --------------------------------------------
def solve_math_problem_with_wolfram(query):
    api_endpoint = "http://api.wolframalpha.com/v2/query"
    params = {'input': query, 'output': 'json', 'appid': WOLFRAM_APP_ID}
    try:
        response = requests.get(api_endpoint, params=params)
        response.raise_for_status()
        data = response.json()
        pods = data.get('queryresult', {}).get('pods', [])
        if pods:
            for pod in pods:
                if pod.get('title', '').lower() in ['result', 'exact result']:
                    for subpod in pod.get('subpods', []):
                        if subpod.get('plaintext'):
                            return subpod['plaintext']
            for pod in pods:
                for subpod in pod.get('subpods', []):
                    if subpod.get('plaintext'):
                        return subpod['plaintext']
        return "No result found."
    except Exception as e:
        return f"Error querying Wolfram Alpha: {e}"

def solve_math_with_groq_llm(query):
    prompt = f"Solve this mathematical problem with detailed steps: {query}"
    try:
        response = groq_client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Groq LLM error: {e}"

# --------------------------------------------
# Free Dictionary API
# --------------------------------------------
def get_definition_from_free_dictionary(query):
    url = f"https://api.dictionaryapi.dev/api/v2/entries/en/{query.lower()}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
        word = data[0]['word']
        definition = data[0]['meanings'][0]['definitions'][0]['definition']
        return word, definition
    except Exception as e:
        return None, f"Free Dictionary API error: {str(e)}"

# --------------------------------------------
# Logging Helper
# --------------------------------------------
def log_decision(query, route):
    logging.info(f"Query: '{query}' | Route: {route}")

# --------------------------------------------
# LLM Answer via LangChain
# --------------------------------------------
def get_llm_answer(query, docs):
    llm = ChatOpenAI(temperature=0)
    chain = load_qa_chain(llm, chain_type="stuff")
    return chain.run(input_documents=docs, question=query)

# --------------------------------------------
# Query Classifiers
# --------------------------------------------
def is_math_query(query):
    return re.search(r'\b(calculate|plus|minus|times|divide|compute|multiplied|divided|add|subtract|\d+|\=|\^|\*|/|-|\+)\b', query.lower())

def is_definition_query(query):
    return re.search(r'\b(define|what is|meaning of|definition of|explain|describe)\b', query.lower())

# --------------------------------------------
# Agent Router
# --------------------------------------------
def agent_router(query):
    response = {}

    if is_math_query(query):
        try:
            solution = solve_math_problem_with_wolfram(query)
            if "No result found" in solution or "Error" in solution:
                solution = solve_math_with_groq_llm(query)
            response['route'] = '🧮 Calculator (Wolfram → Groq Fallback)'
            response['result'] = solution
        except Exception:
            response['route'] = '🧮 Calculator (Groq Only)'
            response['result'] = solve_math_with_groq_llm(query)
        finally:
            log_decision(query, response['route'])

    elif is_definition_query(query):
        try:
            word = re.sub(r'^(define|what is|meaning of|definition of|explain|describe)\s*', '', query.lower()).strip()
            word, definition = get_definition_from_free_dictionary(word)
            if word and definition:
                response['route'] = '📚 Dictionary (Free Dictionary API)'
                response['result'] = f"Definition of {word}: {definition}"
            else:
                response['route'] = '📚 Dictionary (Not Found)'
                response['result'] = "No definition found."
        except Exception as e:
            response['route'] = '📚 Dictionary (Error)'
            response['result'] = f"Definition error: {e}"
        finally:
            log_decision(query, response['route'])

    else:
        try:
            from retriever import retrieve_top_k  # ensure retriever is implemented
            docs = retrieve_top_k(query)
            context = [doc.page_content[:200] for doc in docs]
            answer = get_llm_answer(query, docs)
            response['route'] = '📖 RAG → LLM'
            response['context'] = context
            response['result'] = answer
        except Exception as e:
            response['route'] = '📖 RAG → LLM (Error)'
            response['result'] = f"RAG error: {e}"
        finally:
            log_decision(query, response['route'])

    return response

# --------------------------------------------
# RAG: Document Ingestion & Retrieval
# --------------------------------------------
def chunk_text(text, max_length=300):
    sentences = re.split(r'[.!?]\s+', text)
    chunks, chunk = [], ""
    for sentence in sentences:
        if len(chunk) + len(sentence) < max_length:
            chunk += sentence + '. '
        else:
            chunks.append(chunk.strip())
            chunk = sentence + '. '
    if chunk:
        chunks.append(chunk.strip())
    return chunks

def read_file(file):
    ext = os.path.splitext(file.name)[-1].lower()
    if ext == '.pdf':
        reader = PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif ext == '.docx':
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == '.csv':
        df = pd.read_csv(file)
        return df.to_string(index=False)
    elif ext in ['.txt', '.md']:
        return file.read().decode("utf-8")
    else:
        return ""

def process_documents(files):
    texts, metadatas = [], []
    for file in files:
        content = read_file(file)
        if content:
            chunks = chunk_text(content)
            texts.extend(chunks)
            metadatas.extend([{"source": file.name}] * len(chunks))
    return texts, metadatas

def build_vector_store(texts):
    embeddings = embed_model.encode(texts)
    index = faiss.IndexFlatL2(embeddings[0].shape[0])
    index.add(embeddings)
    return index, embeddings

def retrieve_top_k_local(query, texts, index, k=3):
    query_vec = embed_model.encode([query])
    _, indices = index.search(query_vec, k)
    return [Document(page_content=texts[i]) for i in indices[0]]

